"""
Converte swot_bringdata.md → swot_bringdata.docx
Dependências: python-docx (já instalado)
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


MD_PATH   = "/Users/ramonmoreira/Desktop/bring_data/V2/docs/swot_bringdata.md"
DOCX_PATH = "/Users/ramonmoreira/Desktop/bring_data/V2/docs/swot_bringdata.docx"

# ── Paleta de cores ──────────────────────────────────────────────────────────
C_BLACK  = RGBColor(0x1A, 0x1A, 0x2E)   # quase preto (títulos)
C_ACCENT = RGBColor(0x16, 0x21, 0x3E)   # azul escuro
C_BODY   = RGBColor(0x2E, 0x2E, 0x2E)   # cinza escuro (corpo)
C_META   = RGBColor(0x66, 0x66, 0x66)   # cinza médio (metadados)
C_RULE   = RGBColor(0xCC, 0xCC, 0xCC)   # cinza claro (separadores)

# Mapeamento de prefixo de seção → cor de destaque da barra lateral
SECTION_COLORS = {
    "FORÇAS":      RGBColor(0x27, 0xAE, 0x60),   # verde
    "FRAQUEZAS":   RGBColor(0xE7, 0x4C, 0x3C),   # vermelho
    "OPORTUNIDADES": RGBColor(0x29, 0x80, 0xB9), # azul
    "AMEAÇAS":     RGBColor(0xE6, 0x7E, 0x22),   # laranja
    "SÍNTESE":     RGBColor(0x8E, 0x44, 0xAD),   # roxo
    "CONTEXTO":    RGBColor(0x34, 0x49, 0x5E),   # azul acinzentado
    "FONTES":      RGBColor(0x95, 0xA5, 0xA6),   # cinza
}

CURRENT_SECTION_COLOR = RGBColor(0x34, 0x49, 0x5E)


def set_section_color(heading_text):
    global CURRENT_SECTION_COLOR
    for key, color in SECTION_COLORS.items():
        if key in heading_text.upper():
            CURRENT_SECTION_COLOR = color
            return


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_left_border(paragraph, hex_color):
    """Adiciona barra lateral colorida ao parágrafo."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '6')
    left.set(qn('w:color'), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def rgb_to_hex(color):
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def apply_inline(run_text, paragraph, default_bold=False, default_color=C_BODY):
    """
    Processa markdown inline (bold, italic, inline code) dentro de um parágrafo.
    Adiciona runs com formatação adequada.
    """
    # Padrões inline (ordem importa: code > bold+italic > bold > italic)
    pattern = re.compile(r'(`[^`]+`|\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*|__[^_]+__|_[^_]+_)')
    parts = pattern.split(run_text)

    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        run.font.color.rgb = default_color

        if part.startswith('`') and part.endswith('`'):
            run.text = part[1:-1]
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        elif part.startswith('***') and part.endswith('***'):
            run.text = part[3:-3]
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run.text = part[2:-2]
            run.bold = True
        elif (part.startswith('*') and part.endswith('*')) or \
             (part.startswith('_') and part.endswith('_')):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith('__') and part.endswith('__'):
            run.text = part[2:-2]
            run.bold = True
        else:
            run.text = part
            if default_bold:
                run.bold = True


def parse_table(lines, doc):
    """Cria tabela a partir de linhas markdown."""
    rows = []
    for line in lines:
        if re.match(r'^\s*\|[-| :]+\|\s*$', line):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)

    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'

    # Largura das colunas
    total_width = Inches(6.0)
    col_width = total_width / n_cols
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j >= n_cols:
                break
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            is_header = (i == 0)
            clean = re.sub(r'^\*\*(.+)\*\*$', r'\1', cell_text)
            apply_inline(clean, p, default_bold=is_header,
                         default_color=C_BLACK if is_header else C_BODY)
            if is_header:
                # Fundo cinza suave para cabeçalho
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F2F2F2')
                tcPr.append(shd)

    doc.add_paragraph()


def convert():
    doc = Document()

    # ── Margens do documento ─────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.15)
        section.right_margin  = Inches(1.15)

    # ── Estilo padrão ────────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = C_BODY

    with open(MD_PATH, encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        # ── Linha em branco ──────────────────────────────────────────────────
        if line.strip() == '':
            i += 1
            continue

        # ── Separador --- ────────────────────────────────────────────────────
        if re.match(r'^---+\s*$', line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Tabela ───────────────────────────────────────────────────────────
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].rstrip('\n'))
                i += 1
            parse_table(table_lines, doc)
            continue

        # ── Blockquote ───────────────────────────────────────────────────────
        if line.startswith('>'):
            text = line.lstrip('> ').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            add_left_border(p, rgb_to_hex(CURRENT_SECTION_COLOR))
            apply_inline(text, p, default_color=RGBColor(0x44, 0x44, 0x44))
            for run in p.runs:
                run.font.size = Pt(10)
                run.italic = True
            i += 1
            continue

        # ── Headings ─────────────────────────────────────────────────────────
        h_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if h_match:
            level  = len(h_match.group(1))
            h_text = h_match.group(2).strip()
            set_section_color(h_text)

            p = doc.add_paragraph()

            if level == 1:
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after  = Pt(6)
                run = p.add_run(h_text)
                run.bold = True
                run.font.size = Pt(18)
                run.font.color.rgb = C_BLACK
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            elif level == 2:
                p.paragraph_format.space_before = Pt(16)
                p.paragraph_format.space_after  = Pt(4)
                add_left_border(p, rgb_to_hex(CURRENT_SECTION_COLOR))
                p.paragraph_format.left_indent = Inches(0.15)
                run = p.add_run(h_text.upper())
                run.bold = True
                run.font.size = Pt(13)
                run.font.color.rgb = CURRENT_SECTION_COLOR

            elif level == 3:
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after  = Pt(3)
                run = p.add_run(h_text)
                run.bold = True
                run.font.size = Pt(11.5)
                run.font.color.rgb = CURRENT_SECTION_COLOR

            else:
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after  = Pt(2)
                run = p.add_run(h_text)
                run.bold = True
                run.font.size = Pt(10.5)
                run.font.color.rgb = C_ACCENT

            i += 1
            continue

        # ── Lista com marcador ───────────────────────────────────────────────
        bullet_match = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if bullet_match:
            indent_spaces = len(bullet_match.group(1))
            text = bullet_match.group(2).strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            if indent_spaces >= 2:
                p.paragraph_format.left_indent = Inches(0.4 + indent_spaces * 0.02)
            apply_inline(text, p, default_color=C_BODY)
            i += 1
            continue

        # ── Linha normal ─────────────────────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        apply_inline(line.strip(), p, default_color=C_BODY)
        i += 1

    doc.save(DOCX_PATH)
    print(f"✓ Arquivo salvo em: {DOCX_PATH}")


if __name__ == '__main__':
    convert()
