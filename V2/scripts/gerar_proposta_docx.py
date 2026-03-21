"""
Gera smart_ads_proposta_v2.docx com formatação equivalente ao padrão Google Docs.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import copy

OUTPUT = Path(__file__).parent.parent / "propostas_e_apresentacoes" / "smart_ads_proposta_v2.docx"

# ── cores ──────────────────────────────────────────────────────────────────────
BLACK      = RGBColor(0x1a, 0x1a, 0x1a)
DARK_GRAY  = RGBColor(0x44, 0x44, 0x44)
MID_GRAY   = RGBColor(0x77, 0x77, 0x77)
GREEN      = RGBColor(0x1d, 0x8a, 0x3e)
GREEN_LIGHT= RGBColor(0xe8, 0xf5, 0xec)
YELLOW_BG  = RGBColor(0xff, 0xf8, 0xe1)
BLUE_DARK  = RGBColor(0x0d, 0x47, 0xa1)
WHITE      = RGBColor(0xff, 0xff, 0xff)
ROW_ALT    = RGBColor(0xf5, 0xf5, 0xf5)
HEADER_BG  = RGBColor(0x1a, 0x1a, 0x1a)

def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=False, bottom=False, color="DDDDDD"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, show in [('top', top), ('bottom', bottom), ('left', False), ('right', False)]:
        el = OxmlElement(f'w:{side}')
        if show:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:color'), color)
        else:
            el.set(qn('w:val'), 'none')
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color=BLACK, space_before=18, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    sizes = {1: 22, 2: 16, 3: 13}
    run.font.size = Pt(sizes.get(level, 13))
    run.font.bold = True
    run.font.color.rgb = color
    return p

def add_para(doc, text, bold_parts=None, size=11, color=DARK_GRAY, space_after=8,
             italic=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)

    if bold_parts is None:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.italic = italic
    else:
        # bold_parts: list of (text, bold, color_override)
        for seg_text, seg_bold, seg_color in bold_parts:
            r = p.add_run(seg_text)
            r.font.size = Pt(size)
            r.font.bold = seg_bold
            r.font.color.rgb = seg_color or color
    return p

def add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_callout(doc, text, bg=GREEN_LIGHT, text_color=GREEN, border_color="1d8a3e"):
    """Caixa destacada (garantia, etc)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    cell.width = Inches(6)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = text_color
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_margin_table(doc):
    headers = ["Período", "Margem com ML", "Sem ML (est.)*", "Ganho ML"]
    rows = [
        ["dez/2025  (3 lançamentos)", "+R$ 104.000", "+R$ 80.000",  "+R$ 24.000"],
        ["jan/2026",                  "+R$ 340.000", "+R$ 195.000", "+R$ 145.000"],
        ["fev/2026  (2 lançamentos)", "+R$ 204.000", "+R$ 38.000",  "+R$ 166.000"],
    ]
    subtotal  = ["Subtotal verificado", "+R$ 648.000", "+R$ 313.000", "+R$ 335.000"]
    lf45_row  = ["mar/2026 — 100% ML",   "+R$ 140.000", "+R$ 39.000",  "+R$ 101.000"]
    total_row = ["Total — 4 meses",     "+R$ 788.000", "+R$ 352.000", "+R$ 436.000"]

    n_rows = 1 + len(rows) + 1 + 1 + 1   # header + data + subtotal + lf45 + total
    tbl = doc.add_table(rows=n_rows, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'

    col_widths = [Cm(4.8), Cm(3.8), Cm(3.8), Cm(3.4)]

    def write_row(row_obj, data, bg, text_color=DARK_GRAY, bold=False, size=10):
        for ci, (val, w) in enumerate(zip(data, col_widths)):
            cell = row_obj.cells[ci]
            cell.width = w
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(5 if bold else 4)
            p.paragraph_format.space_after  = Pt(5 if bold else 4)
            p.paragraph_format.left_indent  = Cm(0.2)
            run = p.add_run(val)
            run.font.size  = Pt(size)
            run.font.bold  = bold or (ci == 3)
            run.font.color.rgb = GREEN if ci == 3 else text_color

    # Header
    hdr = tbl.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr.cells[i]
        cell.width = w
        set_cell_bg(cell, HEADER_BG)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        p.paragraph_format.left_indent  = Cm(0.2)
        run = p.add_run(h)
        run.font.size  = Pt(10)
        run.font.bold  = True
        run.font.color.rgb = WHITE

    # Data rows
    for ri, row_data in enumerate(rows):
        bg = ROW_ALT if ri % 2 == 0 else WHITE
        write_row(tbl.rows[ri + 1], row_data, bg)

    # Subtotal row (dark gray background)
    sub_bg = RGBColor(0x37, 0x47, 0x4f)
    sub_row = tbl.rows[1 + len(rows)]
    for ci, (val, w) in enumerate(zip(subtotal, col_widths)):
        cell = sub_row.cells[ci]
        cell.width = w
        set_cell_bg(cell, sub_bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.left_indent  = Cm(0.2)
        run = p.add_run(val)
        run.font.size  = Pt(10)
        run.font.bold  = True
        run.font.color.rgb = WHITE

    # LF45 row (light green — estimated)
    write_row(tbl.rows[1 + len(rows) + 1], lf45_row, GREEN_LIGHT)

    # Grand total row
    tr = tbl.rows[-1]
    for ci, (val, w) in enumerate(zip(total_row, col_widths)):
        cell = tr.cells[ci]
        cell.width = w
        set_cell_bg(cell, RGBColor(0x1d, 0x8a, 0x3e))
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.left_indent  = Cm(0.2)
        run = p.add_run(val)
        run.font.size  = Pt(10)
        run.font.bold  = True
        run.font.color.rgb = WHITE

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_roas_table(doc):
    headers = ["Lançamento", "ROAS ML", "ROAS Controle", "Vantagem ML"]
    rows = [
        ["LF40 — dez/2025", "1,06×", "0,91×", "+17%"],
        ["LF41 — dez/2025", "3,60×", "3,29×", "+10%"],
        ["LF42 — dez/2025", "3,20×", "1,39×", "+130%"],
        ["DEV19 — jan/2026", "3,64×", "1,89×", "+93%"],
        ["LF43 — fev/2026", "3,84×", "1,36×", "+183%"],
        ["LF44 — fev/2026", "4,07×", "1,20×", "+240%"],
    ]

    tbl = doc.add_table(rows=1 + len(rows), cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    col_widths = [Cm(4.2), Cm(2.8), Cm(3.2), Cm(3.0)]

    hdr = tbl.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr.cells[i]
        cell.width = w
        set_cell_bg(cell, HEADER_BG)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Cm(0.2)
        run = p.add_run(h)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WHITE

    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg = ROW_ALT if ri % 2 == 0 else WHITE
        for ci, (val, w) in enumerate(zip(row_data, col_widths)):
            cell = row.cells[ci]
            cell.width = w
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Cm(0.2)
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GRAY
            if ci == 3:
                run.font.bold = True
                run.font.color.rgb = GREEN

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default paragraph font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ── CAPA ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Smart Ads — Proposta de Parceria")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = BLACK

    add_para(doc,
        "Enquanto seus concorrentes otimizam para leads, você pode otimizar para compradores.",
        size=14, color=DARK_GRAY, space_after=6,
        bold_parts=[
            ("Enquanto seus concorrentes otimizam para leads, ", False, DARK_GRAY),
            ("você pode otimizar para compradores.", True, BLACK),
        ])

    add_rule(doc)

    # ── 1. O QUE O MERCADO FAZ ────────────────────────────────────────────────
    add_heading(doc, "O que o mercado faz hoje — e por que funciona mal", level=2,
                space_before=16, space_after=6)

    add_para(doc,
        "A forma mais comum de tentar qualificar leads é o lead scoring por regras: "
        "estudar quem comprou, identificar características em comum, e atribuir pontos "
        "fixos para cada variável. O problema é que esse sistema é linear e cego para interações. "
        "Ele não enxerga que a combinação de três variáveis pode ser 10× mais preditiva do que qualquer "
        "uma isolada. E as regras que funcionavam em março podem não funcionar em setembro — "
        "elas nunca se atualizam sozinhas.",
        space_after=8)

    add_para(doc,
        "Um modelo de Machine Learning treinado nos seus dados resolve exatamente isso. "
        "Ele não tem conhecimento prévio: aprende exclusivamente do histórico do seu negócio. "
        "Encontra padrões não-lineares que nenhuma regra manual conseguiria mapear. "
        "Otimiza uma métrica real — a capacidade de separar compradores de não-compradores nos seus leads, "
        "em tempo real, já validado contra dados que nunca viu. "
        "E o que ele aprende sobre o perfil de comprador do seu negócio não existe em nenhum outro lugar.",
        space_after=8)

    add_rule(doc)

    # ── 2. O QUE O SMART ADS FAZ ─────────────────────────────────────────────
    add_heading(doc, "Como o Smart Ads funciona", level=2, space_before=14, space_after=6)

    add_para(doc,
        "A cada lead que preenche a pesquisa, o Smart Ads atribui um score de propensão à compra "
        "e envia esse sinal ao Meta via CAPI — em menos de 5 minutos. "
        "O Meta aprende que aquele perfil de pessoa converte, passa a encontrá-la com mais eficiência "
        "no leilão e reduz o custo por lead. Com o mesmo orçamento, você chega a mais compradores — "
        "e paga menos por cada um.",
        space_after=8)

    add_para(doc,
        "O modelo usa muito mais do que as respostas da pesquisa: UTMs de origem, "
        "histórico do lead em lançamentos anteriores, validade dos dados de contato, "
        "padrões temporais e, nas próximas versões, dados comportamentais como percentual de vídeo "
        "assistido e interações no WhatsApp. Quanto mais dados disponíveis, mais preciso o sinal.",
        space_after=8)

    add_rule(doc)

    # ── 3. OS RESULTADOS ──────────────────────────────────────────────────────
    add_heading(doc, "Os resultados reais — 4 meses de operação", level=2,
                space_before=14, space_after=6)

    add_para(doc,
        "O sistema está em produção desde dezembro de 2025. "
        "A tabela abaixo compara a margem de contribuição real gerada com ML "
        "contra o que teria sido gerado se o mesmo orçamento fosse aplicado com "
        "a eficiência das campanhas sem ML do mesmo período.",
        space_after=8)

    add_para(doc,
        "Em três meses com grupo Controle rodando simultaneamente, o sistema gerou "
        "R$ 335.000 de margem incremental verificada. "
        "Ao migrar 100% do orçamento para ML em março (LF45), o resultado estimado "
        "pelo baseline histórico soma R$ 436.000 no período completo de 4 meses.",
        bold_parts=[
            ("Em três meses com grupo Controle rodando simultaneamente, o sistema gerou ", False, DARK_GRAY),
            ("R$ 335.000 de margem incremental verificada.", True, BLACK),
            (" Ao migrar 100% do orçamento para ML em março (LF45), o resultado estimado "
             "pelo baseline histórico soma ", False, DARK_GRAY),
            ("R$ 436.000 no período completo de 4 meses.", True, BLACK),
        ],
        space_after=10)

    add_margin_table(doc)

    add_para(doc,
        "* Para dez/2025–fev/2026: margem contrafactual calculada com o ROAS do grupo Controle "
        "que rodou simultaneamente ao ML no mesmo período. "
        "† LF45 (mar/2026): sem grupo Controle; cliente migrou 100% do orçamento para ML. "
        "Contrafactual estimado pela mediana histórica do ROAS Controle dos seis períodos anteriores (1,36×).",
        size=9, color=MID_GRAY, space_after=10, italic=True)

    # ROAS table
    add_heading(doc, "ROAS ML versus Controle — 6 lançamentos consecutivos", level=3,
                space_before=12, space_after=6)

    add_para(doc,
        "Nos seis períodos em que campanhas com e sem ML rodaram simultaneamente, "
        "o ROAS das campanhas ML foi superior em todos — de 10% a 240% acima do Controle. "
        "O CPL das campanhas ML ficou entre 28% e 44% abaixo em todos os períodos.",
        space_after=10)

    add_roas_table(doc)

    add_para(doc,
        "A vantagem não vem de uma taxa de conversão artificialmente alta: "
        "ela vem do custo menor de aquisição do lead. O Meta aprende o sinal "
        "e passa a entregar o mesmo perfil de comprador a um preço menor no leilão. "
        "A partir do terceiro lançamento, o efeito se torna crescente — "
        "o modelo acumula inteligência a cada retreino.",
        space_after=8)

    add_rule(doc)

    # ── 4. SUSTENTABILIDADE ───────────────────────────────────────────────────
    add_heading(doc, "O que garante o resultado no longo prazo", level=2,
                space_before=14, space_after=6)

    add_para(doc,
        "Comportamentos de compra mudam. Um modelo que funciona hoje pode degradar em seis meses "
        "se não for atualizado. O Smart Ads inclui retreino periódico com os dados mais recentes "
        "de conversão do seu negócio — garantindo que o sinal enviado ao Meta continue preciso. "
        "Um painel de monitoramento acompanha diariamente a qualidade dos dados, "
        "a distribuição dos leads e desvios no público. Se algo começa a mudar, "
        "você é avisado antes de afetar o resultado.",
        space_after=8)

    add_rule(doc)

    # ── 5. INVESTIMENTO ───────────────────────────────────────────────────────
    add_heading(doc, "Investimento", level=2, space_before=14, space_after=6)

    add_para(doc,
        "Para montar um sistema equivalente internamente, você precisaria de pelo menos três perfis: "
        "Cientista de Dados, Engenheiro de Dados e Engenheiro de ML. "
        "Só em folha: R$38.000/mês no mínimo — com encargos e benefícios, R$70.000 a R$80.000/mês. "
        "Sem garantia de resultado, sem prazo e com meses de curva de aprendizado antes de qualquer "
        "modelo ir a produção. E ainda assim faltaria o mais difícil: nenhum desses profissionais "
        "chega com sete anos de contexto sobre o mercado de educação online.",
        space_after=12)

    # ── Opção A — Fee Fixo
    add_heading(doc, "Opção A — Fee Fixo", level=3, space_before=8, space_after=6)

    tbl_a = doc.add_table(rows=1, cols=2)
    tbl_a.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_a.style = 'Table Grid'
    for ci, (val, lbl, bg, fc) in enumerate([
        ("R$ 17.500",     "Setup\n(desenvolvimento, integração e\nprimeira versão do modelo)",
         RGBColor(0xf0,0xf4,0xff), BLUE_DARK),
        ("R$ 15.000/mês", "Mensalidade\n(monitoramento, retreino,\nrelatórios e suporte contínuo)",
         RGBColor(0xe8,0xf5,0xec), GREEN),
    ]):
        cell = tbl_a.cell(0, ci)
        cell.width = Cm(8.0)
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.left_indent  = Cm(0.5)
        r1 = p.add_run(val + "\n")
        r1.font.size = Pt(18); r1.font.bold = True; r1.font.color.rgb = fc
        r2 = p.add_run(lbl)
        r2.font.size = Pt(9);  r2.font.color.rgb = MID_GRAY

    add_para(doc, "Previsibilidade total. Indicado para quem já tem clareza do retorno esperado "
             "ou investe acima de R$200k/mês em anúncios.",
             size=10, color=MID_GRAY, space_after=12, italic=True)

    # ── Opção B — Rev Share
    add_heading(doc, "Opção B — Rev Share", level=3, space_before=8, space_after=6)

    tbl_b = doc.add_table(rows=1, cols=2)
    tbl_b.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_b.style = 'Table Grid'
    for ci, (val, lbl, bg, fc) in enumerate([
        ("Sem setup",    "Nenhum investimento inicial.\nO sistema entra em operação\nsem custo de entrada.",
         RGBColor(0xf5,0xf5,0xf5), DARK_GRAY),
        ("20% + mín. R$5k", "Da margem incremental mensal.\nSó paga valor maior quando\no sistema gera resultado de fato.",
         RGBColor(0xe8,0xf5,0xec), GREEN),
    ]):
        cell = tbl_b.cell(0, ci)
        cell.width = Cm(8.0)
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.left_indent  = Cm(0.5)
        r1 = p.add_run(val + "\n")
        r1.font.size = Pt(18); r1.font.bold = True; r1.font.color.rgb = fc
        r2 = p.add_run(lbl)
        r2.font.size = Pt(9);  r2.font.color.rgb = MID_GRAY

    add_para(doc, "Indicado para quem quer começar sem compromisso fixo. "
             "Você paga o mínimo de R$5.000 em meses neutros e 20% da margem incremental nos demais.",
             size=10, color=MID_GRAY, space_after=10, italic=True)

    # Fórmula
    add_heading(doc, "Como é calculada a margem incremental", level=3,
                space_before=8, space_after=4)

    add_para(doc,
        "Margem incremental = margem real do lançamento − o que teria sido gerado "
        "com o mesmo orçamento sem as campanhas de ML.",
        bold_parts=[
            ("Margem incremental = ", True, BLACK),
            ("margem real do lançamento", False, DARK_GRAY),
            (" − o que teria sido gerado com o mesmo orçamento sem as campanhas de ML.", False, DARK_GRAY),
        ], space_after=6)

    # Exemplo numérico
    ex_headers = ["", "Com ML", "Sem ML", ""]
    ex_rows = [
        ["Investimento em anúncios", "R$ 100.000", "R$ 100.000", "—"],
        ["Receita gerada",           "R$ 180.000", "R$ 140.000", ""],
        ["Margem",                   "R$ 80.000",  "R$ 40.000",  ""],
        ["Margem incremental",       "",            "",           "R$ 40.000"],
        ["Rev share (20%)",          "",            "",           "R$ 8.000"],
    ]
    tbl_ex = doc.add_table(rows=1 + len(ex_rows), cols=4)
    tbl_ex.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_ex.style = 'Table Grid'
    ex_widths = [Cm(5.2), Cm(3.0), Cm(3.0), Cm(2.6)]

    hdr_ex = tbl_ex.rows[0]
    for i, (h, w) in enumerate(zip(ex_headers, ex_widths)):
        cell = hdr_ex.cells[i]
        cell.width = w
        set_cell_bg(cell, HEADER_BG)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Cm(0.2)
        run = p.add_run(h)
        run.font.size = Pt(9); run.font.bold = True; run.font.color.rgb = WHITE

    bold_ex_rows = {2, 3, 4}
    for ri, row_data in enumerate(ex_rows):
        row = tbl_ex.rows[ri + 1]
        bg = ROW_ALT if ri % 2 == 0 else WHITE
        for ci, (val, w) in enumerate(zip(row_data, ex_widths)):
            cell = row.cells[ci]
            cell.width = w
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Cm(0.2)
            run = p.add_run(val)
            run.font.size = Pt(9)
            is_bold_row = ri in bold_ex_rows
            run.font.bold = is_bold_row
            run.font.color.rgb = GREEN if (ci == 3 and val) else DARK_GRAY

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_para(doc,
        "Enquanto ML e Controle rodam juntos, o cálculo usa o ROAS do grupo Controle "
        "observado naquele período — auditável diretamente no Meta Ads, sem nenhuma estimativa. "
        "O baseline fixo é estabelecido nos primeiros 3 meses e registrado em contrato, "
        "e passa a ser usado apenas quando o Controle é reduzido ou eliminado.",
        size=10, color=DARK_GRAY, space_after=8)

    # ── Opção C — Poucos lançamentos por ano
    add_heading(doc, "Opção C — Poucos lançamentos por ano (3–4/ano)", level=3,
                space_before=8, space_after=6)

    add_para(doc,
        "Para clientes com 3 ou 4 grandes lançamentos anuais, o modelo padrão de mensalidade "
        "fixa durante os meses inativos não faz sentido econômico. "
        "Duas alternativas foram estruturadas para esse perfil:",
        size=10, color=DARK_GRAY, space_after=8)

    # C1 — Rev Share com mínimo apenas nos meses ativos
    add_heading(doc, "C1 — Rev Share com mínimo apenas nos meses de captação ativa",
                level=3, space_before=4, space_after=4)

    tbl_c1 = doc.add_table(rows=1, cols=2)
    tbl_c1.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_c1.style = 'Table Grid'
    for ci, (val, lbl, bg, fc) in enumerate([
        ("R$ 17.500",        "Setup único\n(desenvolvimento, integração e\nprimeira versão do modelo)",
         RGBColor(0xf0,0xf4,0xff), BLUE_DARK),
        ("20% + mín. R$5k", "Da margem incremental.\nMínimo cobrado apenas nos\nmeses com captação ativa.",
         RGBColor(0xe8,0xf5,0xec), GREEN),
    ]):
        cell = tbl_c1.cell(0, ci)
        cell.width = Cm(8.0)
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.left_indent  = Cm(0.5)
        r1 = p.add_run(val + "\n")
        r1.font.size = Pt(18); r1.font.bold = True; r1.font.color.rgb = fc
        r2 = p.add_run(lbl)
        r2.font.size = Pt(9); r2.font.color.rgb = MID_GRAY

    add_para(doc,
        "Nos meses entre lançamentos, sem captação ativa, não há cobrança. "
        "O sistema permanece monitorado e pronto para o próximo ciclo.",
        size=10, color=MID_GRAY, space_after=12, italic=True)

    # C2 — Fee fixo por lançamento
    add_heading(doc, "C2 — Fee fixo por lançamento realizado",
                level=3, space_before=4, space_after=4)

    tbl_c2 = doc.add_table(rows=1, cols=2)
    tbl_c2.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_c2.style = 'Table Grid'
    for ci, (val, lbl, bg, fc) in enumerate([
        ("R$ 17.500",    "Setup único\n(desenvolvimento, integração e\nprimeira versão do modelo)",
         RGBColor(0xf0,0xf4,0xff), BLUE_DARK),
        ("R$ 9.500\npor lançamento", "Fee fixo por lançamento realizado.\nCobrado no início de cada\ncaptação, independente do resultado.",
         RGBColor(0xe8,0xf5,0xec), GREEN),
    ]):
        cell = tbl_c2.cell(0, ci)
        cell.width = Cm(8.0)
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.left_indent  = Cm(0.5)
        r1 = p.add_run(val + "\n")
        r1.font.size = Pt(16); r1.font.bold = True; r1.font.color.rgb = fc
        r2 = p.add_run(lbl)
        r2.font.size = Pt(9); r2.font.color.rgb = MID_GRAY

    add_para(doc,
        "Previsibilidade total por lançamento. Para 4 lançamentos/ano equivale a R$38.000 anuais "
        "— sem surpresas e sem cobrança nos meses sem captação.",
        size=10, color=MID_GRAY, space_after=10, italic=True)

    add_rule(doc)

    # ── 6. GARANTIA ───────────────────────────────────────────────────────────
    add_heading(doc, "Garantia de performance", level=2, space_before=14, space_after=8)

    # Definição de Controle
    add_heading(doc, "O que é o grupo Controle", level=3, space_before=6, space_after=4)
    add_para(doc,
        "Controle são as campanhas de captação do cliente sem sinal de ML, rodando "
        "simultaneamente durante os primeiros meses. O ROAS médio dessas campanhas, "
        "calculado nos 3 primeiros meses qualificáveis, é fixado em contrato como baseline "
        "de referência para toda a parceria. São excluídos do cálculo: períodos sazonais atípicos "
        "(novembro/Black Friday e janeiro), lançamentos com produto ou ticket diferente do padrão, "
        "promoções fora do calendário habitual e lançamentos-teste com volume insuficiente de leads.",
        space_after=10)

    # Caixa de garantia
    add_callout(doc,
        "Após 3 meses de operação: se o ROAS das campanhas ML ficar abaixo do baseline Controle "
        "por 2 meses consecutivos, você pode suspender o pagamento até a recuperação "
        "ou cancelar o contrato sem multa.",
        bg=GREEN_LIGHT, text_color=GREEN)

    add_para(doc,
        "A garantia compara ML versus baseline Controle — não versus o resultado absoluto do lançamento. "
        "Em períodos adversos de mercado, o ML pode seguir cumprindo a garantia mesmo quando "
        "o lançamento fecha no negativo: ele protege a margem, não elimina o risco de mercado.",
        space_after=10)

    add_para(doc,
        "Contrato com adesão mínima de 6 meses. Após isso, cancelamento sem multa com 30 dias de aviso. "
        "No cancelamento antecipado no modelo Fee Fixo, aplica-se taxa de desligamento equivalente ao setup — "
        "zerada se o sistema não estiver cumprindo a garantia de performance ou em caso de "
        "encerramento formal e documentado da relação de coprodução. "
        "No modelo Rev Share não há taxa de desligamento.",
        space_after=8)

    add_rule(doc)

    # ── 6b. PERGUNTAS FREQUENTES ──────────────────────────────────────────────
    add_heading(doc, "Perguntas frequentes", level=2, space_before=14, space_after=8)

    faqs = [
        ("E se a Meta mudar o algoritmo?",
         "Mudanças de algoritmo afetam todas as campanhas igualmente. O sinal ML é uma vantagem "
         "relativa — e as atualizações recentes do Meta favoreceram quem envia eventos CAPI de qualidade. "
         "A integração via CAPI é uma funcionalidade oficial recomendada pela própria Meta."),
        ("E se o lançamento tiver prejuízo independente do ML?",
         "A garantia compara ML versus baseline — não versus lucro absoluto. "
         "Em períodos adversos, o ML protege a margem: nos dados reais, o sistema entregou ROAS "
         "superior ao Controle mesmo em lançamentos que fecharam no negativo."),
        ("E se eu mudar o produto, o ticket ou a pesquisa?",
         "Mudanças relevantes no produto, oferta ou ticket acionam um novo ciclo de retreino. "
         "O modelo se adapta — desde que a mudança seja comunicada para incorporar "
         "os novos dados de conversão no próximo ciclo de atualização."),
        ("E se minha conta de anúncios for banida?",
         "O CAPI é uma integração oficial recomendada pela Meta. "
         "Eventual banimento de conta por outros motivos é força maior, "
         "fora do escopo de qualquer prestador de serviço de tráfego."),
    ]

    for question, answer in faqs:
        add_para(doc, question, bold_parts=[(question, True, BLACK)], space_after=3)
        add_para(doc, answer, size=10, color=DARK_GRAY, space_after=10, indent=True)

    add_rule(doc)

    # ── 7. POR QUE NÃO FAZER INTERNAMENTE ────────────────────────────────────
    add_heading(doc, "Por que não fazer isso internamente com meu time?",
                level=2, space_before=14, space_after=6)

    add_para(doc,
        "Pode fazer. Mas considere o custo real: além dos R$70–80k/mês em folha, "
        "você enfrentaria meses de desenvolvimento sem garantia de resultado, "
        "e ainda precisaria de alguém que entenda onde o sinal de qualidade se esconde "
        "nos dados de uma operação de lançamentos — o que vale, o que distorce, "
        "o que o Meta interpreta bem e o que ele ignora. "
        "Esse conhecimento não está em nenhum currículo. "
        "É o que separa um modelo que funciona no laboratório de um modelo que funciona "
        "com orçamento real em jogo.",
        space_after=10)

    add_para(doc,
        "O Smart Ads oferece a alternativa: sistema funcionando desde o primeiro dia, "
        "sem meses de desenvolvimento às suas custas, "
        "com contexto acumulado de mais de 120 lançamentos executados.",
        space_after=8)

    add_rule(doc)

    # ── 8. BIO ────────────────────────────────────────────────────────────────
    add_heading(doc, "Sobre o desenvolvedor", level=2, space_before=14, space_after=6)

    add_para(doc,
        "Profissional com 7 anos de atuação no mercado de educação online, "
        "tendo executado mais de 120 lançamentos de autoria própria e gerenciado "
        "operações de múltiplos 7 e 8 dígitos. "
        "Posteriormente se especializou em Machine Learning e MLOps nas seguintes instituições:",
        space_after=6)

    for curso in [
        "Stanford University — Machine Learning",
        "University of Michigan — Python For Everyone",
        "DeepLearning.ai — Machine Learning in Production",
        "Google — Machine Learning Engineer Certificate",
        "MLOps Community — First Stack MLOps",
        "IBM — Databases and SQL for Data Science with Python",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(curso)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GRAY

    doc.add_paragraph()

    add_para(doc,
        '"In God we trust, all others bring data."  — W. Edwards Deming',
        size=10, color=MID_GRAY, italic=True, space_after=4)

    doc.save(OUTPUT)
    print(f"Salvo em: {OUTPUT}")


if __name__ == "__main__":
    build()
