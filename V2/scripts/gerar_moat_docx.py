"""
Gera moat_valor_negocio.docx com a seção reescrita de Moat / Valor do Negócio.
"""

from docx import Document
from docx.shared import Pt
from pathlib import Path

OUTPUT = Path(__file__).parent.parent / "outputs" / "moat_valor_negocio.docx"

def h(doc, text, level=1):
    doc.add_heading(text, level=level)

def b(doc, text):
    doc.add_paragraph(text, style="List Bullet")

def p(doc, text):
    doc.add_paragraph(text)

def main():
    doc = Document()

    h(doc, "Moat / Valor do Negócio", 1)

    # 1. Barreiras
    h(doc, "Barreiras de entrada (por que é difícil de copiar)", 2)
    b(doc, "Funciona para perpétuo com produto médio/alto ticket — não é só lançamento.")
    b(doc, "É diferente de análise de dados por LLM — LLM não otimiza campanhas em tempo real.")
    b(doc, "Barreira técnica: não basta ter o modelo, precisa de MLOps para servir, monitorar e retreinar. Um erro aqui tem custo direto em verba de anúncios.")
    b(doc, "Barreira de contexto de negócio: o modelo aprende padrões específicos do nicho, do produto, do público. Não é transferível de um concorrente para outro sem dados.")
    b(doc, "Mesmo com agentes de IA criando pipelines de ML, a disciplina de MLOps (drift, retreino, paridade treino/produção) ainda requer julgamento humano especializado.")
    b(doc, "Data flywheel com transfer learning: cada lançamento gera dados que melhoram o modelo do próximo. Quanto mais tempo usando, maior a vantagem acumulada.")

    # 2. LeadQualified
    h(doc, "O que o sistema envia à Meta (LeadQualified — o coração do produto)", 2)
    p(doc, "Este é o ponto central da proposta de valor. Em vez de esperar 21 dias por uma compra real, enviamos um sinal de qualidade em ~5 minutos após o lead chegar.")

    h(doc, "Velocidade do sinal", 3)
    b(doc, "Com ML: sinal de qualidade em 5 min após o lead")
    b(doc, "Sem ML: sinal real só após compra, 7–21 dias depois")

    h(doc, "Volume de eventos", 3)
    b(doc, "Com ML: milhares de eventos qualificados por lançamento")
    b(doc, "Sem ML: dezenas ou centenas de compras reais em 21 dias")

    h(doc, "Eliminação do cold start", 3)
    b(doc, "Com ML: Meta atinge mínimo de otimização na primeira semana de captação")
    b(doc, "Sem ML: algumas conversões após 3–4 semanas — Meta nunca sai do cold start durante o lançamento")

    h(doc, "Qualidade do sinal (o diferencial técnico)", 3)
    p(doc, "Nosso modelo usa features que a Meta não tem acesso: respostas da pesquisa de pré-inscrição, sinais de risco/inadimplência da TMB, combinação de UTMs, dados demográficos. O resultado do modelo — o score — é codificado como o valor do evento (D10=R$1.000, D9=R$900...). A Meta recebe um número de qualidade calibrado, não features brutas. Diferente do que a Meta faria sozinha (pegar tudo e jogar num bolo), nosso modelo seleciona as features com pesos otimizados para o produto específico do cliente.")

    h(doc, "Controle e transparência", 3)
    b(doc, "Você sabe exatamente qual feature está influenciando o score")
    b(doc, "Meta é caixa preta: não explica por que priorizou um lead")

    h(doc, "Calibração contínua", 3)
    b(doc, "O modelo é retreinado com dados reais de conversão de cada lançamento")
    b(doc, "O sinal melhora a cada ciclo")

    h(doc, "Portabilidade multi-plataforma", 3)
    b(doc, "O mesmo modelo alimenta Google Ads, TikTok, LinkedIn com o mesmo score")
    b(doc, "Não é necessário reconfigurar em cada plataforma")
    b(doc, "Resultado: 1 modelo → N canais → otimização cruzada (não só dentro da Meta, mas entre canais)")

    # 3. Purchase events
    h(doc, "Fechamento do funil: Purchase events", 2)
    p(doc, "Após o fechamento do carrinho, enviamos os compradores reais à Meta com seus cookies (FBP/FBC) e o valor real da venda. Isso completa o funil no pixel.")

    h(doc, "O que isso adiciona", 3)
    b(doc, "Melhora a qualidade do evento \"Comprar\" no pixel (hoje 4.4/10 — sem CAPI server-side)")
    b(doc, "Cria lookalike audiences de compradores reais com valor para os próximos lançamentos")
    b(doc, "Conecta o loop completo: lead → lead qualificado → compra real — tudo rastreado server-side")

    h(doc, "O que não fazemos aqui", 3)
    p(doc, "Não enviamos features de pesquisa no Purchase. A Meta não usa custom data arbitrário para otimização — o que importa é a identidade (email + cookies) e o valor. A Meta faz o match com seu próprio grafo interno.")

    h(doc, "Enriquecimento retroativo", 3)
    p(doc, "Temos FBP/FBC de todos os leads desde janeiro de 2025 (backup Cloud SQL + Railway). Isso permite enviar Purchase events com qualidade máxima para os 9 lançamentos anteriores, enriquecendo o histórico do pixel imediatamente.")

    # 4. Inadimplência
    h(doc, "Inadimplência como sinal negativo", 2)
    p(doc, "A integração com TMB permite que o modelo otimize não só por \"vai comprar?\" mas por \"vai comprar E não cancelar/inadimplir?\". Isso é único — nenhuma plataforma de anúncios tem acesso a dados de inadimplência do produto do cliente.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Gerado: {OUTPUT}")

if __name__ == "__main__":
    main()
